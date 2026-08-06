# -*- coding: utf-8 -*-

# SPDX-License-Identifier: GPL-3.0-or-later
#
# CODRUG – Computational Drug Discovery Platform
# Copyright (C) 2024–2026 Moisés Maia
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.

"""Final report generator for CODRUG (STEP 6 "Generate Final Report", .docx).

Mirrors the structure of CODOC's MODULES/module_report.py: pulls its data from the unified
per-job JSON (job_dir/<job_name>.json, written incrementally as STEP 1-6 buttons run) plus the
result files each step already writes to disk (CSVs, USI session JSONs, plots). Kept free of
PyQt so it can be unit-tested and reused without a running GUI.

Each of the six report sections (STEP 1, STEP 2, STEP 3, STEP 4, STEP 5, STEP 6) is built
only from what the job's JSON actually recorded as executed - a step nobody ran is skipped
silently rather than shown as "not run", matching the source methodology template's own
instruction (METODOLOGIA_EXEMPLO.docx).

Internal function/variable names below (_add_step4_section, step6, step7, step8, ...) still use
the pre-renumbering step numbers - only the text actually shown to the user (section bar titles,
progress messages) was updated to match the app's current tab numbering. Same convention used
throughout CODRUG.py."""

from __future__ import annotations

import glob
import io
import json
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Callable, Optional

try:
    from rdkit import Chem
    from rdkit.Chem import AllChem
    from rdkit.Chem.Draw import rdMolDraw2D
except Exception:
    Chem = None
    AllChem = None
    rdMolDraw2D = None

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
    _DOCX_AVAILABLE = True
except Exception:
    _DOCX_AVAILABLE = False

try:
    from PIL import Image
    _PIL_AVAILABLE = True
except Exception:
    _PIL_AVAILABLE = False

COLOR_TITLE_BAR = "BDD6EE"
COLOR_SECTION_BAR = "DAE3F3"
COLOR_ACCENT_BAR = "E2EFDA"

# Fixed introductory section (methodology template's "RELATÓRIO FINAL CODRUG: Introdução"):
# unlike every other section below, this text never changes from job to job, so it's kept
# verbatim here in both languages rather than pulled from the job's JSON. Portuguese is the
# template's own (original) language and is used for every idioma other than "en".
_INTRO_TEXTS = {
    "title": {
        "pt": "RELATÓRIO FINAL CODRUG: Introdução",
        "en": "CODRUG FINAL REPORT: Introduction",
    },
    "paragraph_1": {
        "pt": (
            "O Design de Fármacos Assistido por Computador (CADD) pode envolver duas abordagens "
            "principais: o Design de Fármacos Baseado na Estrutura (SBDD), que utiliza dados "
            "estruturais de biomacromoléculas-alvo, e o Design de Fármacos Baseado no Ligante "
            "(LBDD), que se baseia exclusivamente em dados estruturais e de bioatividade de "
            "pequenas moléculas. Em suma, a primeira abordagem permite classificar compostos com "
            "base em sua afinidade prevista pelos alvos, enquanto a segunda possibilita "
            "categorizar ou ordenar compostos de acordo com sua classe ou bioatividade prevista. "
            "Quando aplicadas a grandes bancos de dados de estruturas de compostos, ambas as "
            "abordagens constituem a triagem virtual de alto rendimento (Virtual HTS). Por fim, "
            "são geradas listas ordenadas; após uma análise de consenso e avaliação in silico das "
            "propriedades farmacocinéticas e toxicológicas (ADMET), essas listas apontam os "
            "compostos mais promissores (Hits). Esses Hits seguem, então, para as etapas "
            "subsequentes de bioensaios pré-clínicos e ensaios clínicos no processo de descoberta "
            "de fármacos."
        ),
        "en": (
            "Computer-Aided Drug Design (CADD) can involve two main approaches: Structure-Based "
            "Drug Design (SBDD), which uses structural data from target biomacromolecules, and "
            "Ligand-Based Drug Design (LBDD), which relies exclusively on structural and "
            "bioactivity data from small molecules. In short, the first approach allows compounds "
            "to be classified based on their predicted affinity for the targets, while the second "
            "makes it possible to categorize or rank compounds according to their predicted class "
            "or bioactivity. When applied to large databases of compound structures, both "
            "approaches constitute high-throughput virtual screening (Virtual HTS). Ultimately, "
            "ranked lists are generated; after a consensus analysis and in silico evaluation of "
            "pharmacokinetic and toxicological properties (ADMET), these lists point to the most "
            "promising compounds (Hits). These Hits then proceed to the subsequent stages of "
            "preclinical bioassays and clinical trials in the drug discovery process."
        ),
    },
    "paragraph_2": {
        "pt": (
            "O CODRUG é um software de código aberto que aplica modelos de aprendizado de máquina "
            "(ML) dentro de uma estrutura LBDD-QSAR. Ele está registrado no Instituto Nacional da "
            "Propriedade Industrial (INPI), no Brasil, sob o número 512025006703-8, e hospedado no "
            "GitHub (https://github.com/moimaian/CODRUG), onde está disponível publicamente sob a "
            "licença GPL-3.0. O CODRUG faz parte de uma plataforma de desenvolvimento integrada às "
            "ferramentas de software CODOC, CODYN e CODEEP, seguindo o fluxo de trabalho ilustrado "
            "na Figura 1."
        ),
        "en": (
            "CODRUG is an open-source software that applies machine learning (ML) models within an "
            "LBDD-QSAR framework. It is registered with the National Institute of Industrial "
            "Property (INPI) in Brazil under number 512025006703-8, and hosted on GitHub "
            "(https://github.com/moimaian/CODRUG), where it is publicly available under the "
            "GPL-3.0 license. CODRUG is part of an integrated development platform alongside the "
            "CODOC, CODYN and CODEEP software tools, following the workflow illustrated in Figure 1."
        ),
    },
    "figure_1_caption": {
        "pt": (
            "Figura 1. Fluxo de trabalho integrado que combina abordagens baseadas em ligantes "
            "(QSAR/ML – CODRUG) e abordagens baseadas na estrutura (Docagem Molecular e Dinâmica "
            "Molecular – CODOC e CODYN), juntamente com o preparo ou geração das bases de dados "
            "(Método De Novo – CODEEP)."
        ),
        "en": (
            "Figure 1. Integrated workflow combining ligand-based approaches (QSAR/ML – CODRUG) "
            "and structure-based approaches (Molecular Docking and Molecular Dynamics – CODOC and "
            "CODYN), together with database preparation or generation (De Novo Method – CODEEP)."
        ),
    },
}


def require_docx() -> None:
    if not _DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not available in the current environment. Install python-docx to generate reports.")


def require_rdkit() -> None:
    if Chem is None or rdMolDraw2D is None:
        raise RuntimeError("RDKit is not available in the current environment. Install RDKit to generate reports.")


# --------------------------------------------------------------------------------------
# Job state loading
# --------------------------------------------------------------------------------------

def load_job_settings(job_dir: str) -> dict[str, Any]:
    """Read the unified per-job JSON (job_dir/<job_name>.json), falling back to the legacy
    STEP-1-only dataset_preparation.json for jobs created before it existed."""
    job_dir = str(job_dir)
    job_name = os.path.basename(os.path.normpath(job_dir))
    for path in (
        os.path.join(job_dir, f"{job_name}.json"),
        os.path.join(job_dir, "dataset_preparation.json"),
    ):
        if not os.path.isfile(path):
            continue
        try:
            with open(path, "r", encoding="utf-8") as handle:
                payload = json.load(handle)
        except (OSError, json.JSONDecodeError):
            continue
        if isinstance(payload, dict):
            return payload
    return {}


def _latest_glob(pattern: str) -> Optional[str]:
    matches = glob.glob(pattern)
    if not matches:
        return None
    return max(matches, key=os.path.getmtime)


def _read_csv(path: Optional[str], **kwargs):
    if not path or not os.path.isfile(path):
        return None
    import pandas as pd
    try:
        return pd.read_csv(path, **kwargs)
    except Exception:
        return None


def _row_count(path: Optional[str]) -> Optional[int]:
    df = _read_csv(path)
    return None if df is None else len(df)


# --------------------------------------------------------------------------------------
# 2D structure rendering (ported from CODOC's module_report.molecule_image_bytes)
# --------------------------------------------------------------------------------------

def molecule_image_bytes(smiles: str, size: tuple[int, int] = (220, 160)) -> Optional[io.BytesIO]:
    if Chem is None or rdMolDraw2D is None or not smiles:
        return None
    try:
        mol = Chem.MolFromSmiles(smiles)
        if mol is None:
            return None
        AllChem.Compute2DCoords(mol)
        drawer = rdMolDraw2D.MolDraw2DCairo(*size)
        drawer.drawOptions().padding = 0.15
        drawer.DrawMolecule(mol)
        drawer.FinishDrawing()
        buffer = io.BytesIO(drawer.GetDrawingText())
        buffer.seek(0)
        return buffer
    except Exception:
        return None


def _composite_2x2_image(paths: list[str], out_path: str) -> Optional[str]:
    """Merges up to 4 existing PNG files into a single 2x2 panel image (top-left, top-right,
    bottom-left, bottom-right), matching the "single figure with 4 components" instructions in
    the methodology template. Silently skipped (returns None) if Pillow is unavailable or fewer
    than 2 usable PNGs are found - SVG-only charts are not converted (no raster conversion
    dependency is assumed to be installed)."""
    if not _PIL_AVAILABLE:
        return None
    pngs = [p for p in paths if p and os.path.isfile(p) and p.lower().endswith(".png")]
    if len(pngs) < 2:
        return None
    try:
        images = [Image.open(p).convert("RGB") for p in pngs[:4]]
        cell_w = max(im.width for im in images)
        cell_h = max(im.height for im in images)
        cols = 2 if len(images) > 1 else 1
        rows = 2 if len(images) > 2 else 1
        canvas = Image.new("RGB", (cell_w * cols, cell_h * rows), "white")
        for idx, im in enumerate(images):
            x = (idx % cols) * cell_w
            y = (idx // cols) * cell_h
            canvas.paste(im, (x, y))
        os.makedirs(os.path.dirname(out_path), exist_ok=True)
        canvas.save(out_path)
        return out_path
    except Exception:
        return None


# --------------------------------------------------------------------------------------
# docx helpers (same pattern as CODOC's module_report.py)
# --------------------------------------------------------------------------------------

def _shade_cell(cell: Any, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_cell_width(cell: Any, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)


def _content_width_cm(document: Any) -> float:
    section = document.sections[0]
    return (section.page_width - section.left_margin - section.right_margin) / 360000


def _bar(document: Any, text: str, color_hex: str) -> None:
    width = _content_width_cm(document)
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_width(cell, width)
    _shade_cell(cell, color_hex)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = True
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def _field_line(document: Any, label: str, value: Any) -> None:
    if value in (None, "", "None"):
        return
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    run_label = paragraph.add_run(f"{label}: ")
    run_label.bold = True
    paragraph.add_run(str(value))


PARAGRAPH_FIRST_LINE_INDENT = Cm(1.25)


def _para(document: Any) -> Any:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = PARAGRAPH_FIRST_LINE_INDENT
    return paragraph


def _add(paragraph: Any, text: str, bold: bool = False) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    run.bold = bold


def _join_bold_list(add: Callable[[str, bool], None], items: list[str]) -> None:
    items = [str(i) for i in items if str(i).strip()]
    for index, item in enumerate(items):
        if index > 0:
            add(", " if index < len(items) - 1 else " and ", False)
        add(item, True)


def _join_bold_list_bi(add: Callable[[str, bool], None], items: list[str], lang: str) -> None:
    """Same as _join_bold_list, but with a language-aware final conjunction (" and " / " e ")."""
    items = [str(i) for i in items if str(i).strip()]
    conjunction = " e " if lang == "pt" else " and "
    for index, item in enumerate(items):
        if index > 0:
            add(", " if index < len(items) - 1 else conjunction, False)
        add(item, True)


def _fmt_num(value: Any) -> str:
    try:
        return f"{float(value):g}"
    except (TypeError, ValueError):
        return str(value)


def _add_image(document: Any, path: Optional[str], width_cm: float = 14.0) -> bool:
    if not path or not os.path.isfile(path):
        return False
    try:
        document.add_picture(path, width=Cm(width_cm))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    except Exception:
        return False


def _add_caption(document: Any, text: str) -> Any:
    """Figure/Table caption paragraph, placed BEFORE the figure/table it describes (call this
    first, then _add_image()/_add_static_table()/etc.). The leading label ("Figura 1." /
    "Figure 1." / "Tabela 1." / "Table 1.", up to and including the first ". ") is bold, the
    rest of the text is plain - no italic anywhere in a caption. Justified (not centered), small
    font, no first-line indent (captions aren't body paragraphs)."""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    split_at = text.find(". ")
    if split_at == -1:
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        return paragraph
    label_run = paragraph.add_run(text[:split_at + 1])
    label_run.bold = True
    label_run.font.size = Pt(9)
    rest_run = paragraph.add_run(text[split_at + 1:])
    rest_run.font.size = Pt(9)
    return paragraph


def _add_dataframe_table(document: Any, df, max_rows: int = 15, max_cols: int = 8) -> None:
    if df is None or df.empty:
        return
    cols = list(df.columns)[:max_cols]
    rows = df[cols].head(max_rows)
    width = _content_width_cm(document)
    col_width = width / max(len(cols), 1)
    table = document.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col_idx, name in enumerate(cols):
        cell = table.cell(0, col_idx)
        _set_cell_width(cell, col_width)
        _shade_cell(cell, COLOR_ACCENT_BAR)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].add_run(str(name)).bold = True
    for _, record in rows.iterrows():
        cells = table.add_row().cells
        for col_idx, name in enumerate(cols):
            _set_cell_width(cells[col_idx], col_width)
            cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            value = record[name]
            text = f"{value:.4g}" if isinstance(value, float) else str(value)
            cells[col_idx].paragraphs[0].add_run(text)
    document.add_paragraph()


def _add_static_table(
    document: Any,
    header: list[str],
    rows: list[tuple[str, Any]],
    col_ratios: Optional[list[float]] = None,
) -> None:
    """Renders a fixed reference table (e.g. STEP 1's Table 1 dataset list) from (row_type,
    payload) tuples: row_type "section" spans a single bold/shaded label across every column (a
    merged section-header row); row_type "data" is a normal row of per-column cell text (a cell's
    text may contain "\\n" for multiple lines, e.g. more than one reference link in the same
    cell)."""
    n_cols = len(header)
    width = _content_width_cm(document)
    ratios = col_ratios or [1.0 / n_cols] * n_cols
    col_widths = [width * r for r in ratios]

    table = document.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col_idx, name in enumerate(header):
        cell = table.cell(0, col_idx)
        _set_cell_width(cell, col_widths[col_idx])
        _shade_cell(cell, COLOR_ACCENT_BAR)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].add_run(name).bold = True

    for row_type, payload in rows:
        cells = table.add_row().cells
        if row_type == "section":
            merged = cells[0].merge(cells[-1])
            _set_cell_width(merged, width)
            _shade_cell(merged, COLOR_SECTION_BAR)
            merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            merged.paragraphs[0].add_run(payload).bold = True
            continue
        for col_idx, text in enumerate(payload):
            cell = cells[col_idx]
            _set_cell_width(cell, col_widths[col_idx])
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            lines = str(text).split("\n")
            cell.paragraphs[0].add_run(lines[0])
            for extra_line in lines[1:]:
                extra_para = cell.add_paragraph(extra_line)
                extra_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()


# --------------------------------------------------------------------------------------
# Section 0: RELATÓRIO FINAL CODRUG: Introdução (fixed intro text + Figure 1, always first)
# --------------------------------------------------------------------------------------

def _add_intro_section(document: Any, app_dir: str, idioma: str = "pt") -> bool:
    """Fixed introductory topic from the methodology template - CADD/LBDD-QSAR background, the
    CODRUG description, and Figure 1 (the CODRUG/CODOC/CODYN/CODEEP integrated workflow diagram,
    shipped at BASE/workflow.png). Unlike every other section, this one doesn't depend on the
    job's JSON at all - it's the same for every job, translated to English only when idioma ==
    "en" (Portuguese, the template's own language, is used otherwise)."""
    lang = "en" if idioma == "en" else "pt"

    _bar(document, _INTRO_TEXTS["title"][lang], COLOR_TITLE_BAR)
    p1 = _para(document)
    p1.add_run(_INTRO_TEXTS["paragraph_1"][lang])

    p2 = _para(document)
    p2.add_run(_INTRO_TEXTS["paragraph_2"][lang])
    document.add_paragraph()

    workflow_path = os.path.join(app_dir, "BASE", "workflow.png")
    if os.path.isfile(workflow_path):
        _add_caption(document, _INTRO_TEXTS["figure_1_caption"][lang])
        _add_image(document, workflow_path)

    document.add_page_break()
    return True


# --------------------------------------------------------------------------------------
# Section 1: STEP 1 - Dataset Preparation
# --------------------------------------------------------------------------------------

_STEP1_TEXTS = {
    "intro_pt": (
        "As bases de dados foram divididas em dois tipos de acordo com suas finalidades. O "
        "primeiro tipo é denominado dataset interno e será utilizado para triagem dos modelos "
        "supervisionados de ML, seja de classificação ou regressão. Seus dados são rotulados por "
        "valores de bioatividade (variável dependente ou valores de Y). O segundo tipo são os "
        "datasets externos que não estão rotulados, pois não sabemos seus valores de bioatividade "
        "experimental, mas apenas seus dados de estrutura através do canonical smiles, que "
        "permitirá a obtenção dos descritores (features, variáveis independentes ou valores de X). "
        "Os datasets externos serão utilizados na aplicação dos melhores modelos de ML para "
        "predição de suas bioatividades, a fim de obtermos listas de ranqueamento dos compostos."
    ),
    "intro_en": (
        "The databases were divided into two types according to their purpose. The first type is "
        "called the internal dataset and will be used to screen the supervised ML models, whether "
        "classification or regression. Its data is labeled with bioactivity values (dependent "
        "variable, or Y values). The second type are the external datasets, which are not "
        "labeled since their experimental bioactivity values are unknown - only their structural "
        "data, via canonical SMILES, is available, allowing the descriptors to be obtained "
        "(features, independent variables, or X values). The external datasets will be used to "
        "apply the best ML models to predict their bioactivities, in order to obtain compound "
        "ranking lists."
    ),
    "external_criteria_pt": (
        "Os critérios de inclusão para a construção dos conjuntos de dados externos basearam-se na "
        "disponibilidade das estruturas químicas em formato SMILES (.smi) — ou em formatos "
        "passíveis de conversão —, priorizando bases de dados de produtos naturais e componentes "
        "da biodiversidade global ou regional. Adicionalmente, foram incluídos fármacos aprovados "
        "pela Food and Drug Administration (FDA) e substâncias em fases clínicas e pré-clínicas de "
        "desenvolvimento obtidos aplicando filtros na base de dados Zinc15 (Tabela 1). A inclusão "
        "desta última classe viabiliza a estratégia de drug repurposing, uma abordagem vantajosa "
        "devido à maior acessibilidade sintética, perfis de drogabilidade favoráveis e processos "
        "regulatórios simplificados, decorrentes da prévia disponibilidade de dados de toxicidade. "
        "Dessa forma, buscou-se mapear e abranger um espaço químico amplo e diversificado para as "
        "análises de QSAR."
    ),
    "external_criteria_en": (
        "The inclusion criteria for building the external datasets were based on the availability "
        "of chemical structures in SMILES format (.smi) - or in convertible formats -, "
        "prioritizing databases of natural products and global or regional biodiversity "
        "components. Additionally, drugs approved by the Food and Drug Administration (FDA) and "
        "substances in clinical and preclinical development stages were included, obtained by "
        "applying filters to the Zinc15 database (Table 1). The inclusion of this latter class "
        "enables a drug repurposing strategy, an advantageous approach due to greater synthetic "
        "accessibility, favorable druggability profiles, and simplified regulatory processes, "
        "resulting from the prior availability of toxicity data. In this way, the aim was to map "
        "and cover a broad and diverse chemical space for the QSAR analyses."
    ),
    "table1_caption_pt": (
        "Tabela 1. Bases de dados de estruturas e bioatividades de produtos naturais e fármacos "
        "para reposicionamento usadas para gerar os datasets interno e externo."
    ),
    "table1_caption_en": (
        "Table 1. Structure and bioactivity databases of natural products and repositioning "
        "drugs used to build the internal and external datasets."
    ),
    "table1_footnote_pt": (
        "*Valores originais antes da aplicação de filtros que eliminam compostos repetidos, com "
        "dados de estrutura (ou bioatividade) vazios ou nulos."
    ),
    "table1_footnote_en": (
        "*Original values before applying filters that remove duplicate compounds and compounds "
        "with empty or null structure (or bioactivity) data."
    ),
    "cell_line": {"pt": "Linhagem Celular", "en": "Cell line"},
    "cell_chembl_id": {"pt": "ID ChEMBL da Célula", "en": "Cell ChEMBL ID"},
}

# Table 1: fixed reference list of structure/bioactivity databases used to build the internal and
# external datasets - standard for now, doesn't change per job. Column headers and dataset/
# section labels are kept in English, matching the methodology template's own choice (these read
# as software/database proper nouns rather than prose to translate).
_TABLE_1_HEADER = ["Dataset name", "Number of compounds*", "Database Link", "Reference Link"]
_TABLE_1_ROWS = [
    ("section", "For Internal dataset (training and validation):"),
    ("data", ["CHEMBL", "19614", "https://www.ebi.ac.uk/chembl/explore/targets/", "https://doi.org/10.1093/nar/gkad1004"]),
    ("section", "For External dataset (Prediction and Ranking):"),
    ("section", "Natural Products Databases"),
    ("data", ["Coconut", "637686", "https://coconut.naturalproducts.net/download", "https://doi.org/10.1186/s13321-020-00478-9"]),
    ("data", ["NuBBE", "1787", "https://nubbe.iq.unesp.br/portal/nubbe-search.html", "https://doi.org/10.1038/s41598-017-07451-x\nhttps://doi.org/10.1021/acs.jcim.8b00619"]),
    ("data", ["South Africa", "814", "https://phabidb.vm.uni-freiburg.de/anpdb/", "https://doi.org/10.1186/s13321-021-00514-2"]),
    ("data", ["North and East Africa", "4880", "https://phabidb.vm.uni-freiburg.de/anpdb/", "http://dx.doi.org/10.1021/acs.jnatprod.7b00283\nhttps://doi.org/10.1002/minf.202000163"]),
    ("data", ["CHEMBL", "8782", "https://www.ebi.ac.uk/chembl/explore/compounds/", "https://doi.org/10.1093/nar/gkad1004"]),
    ("data", ["LIFECHEM", "11234", "https://lifechemicals.com/", "https://lifechemicals.com/the-company/life-chemicals-team"]),
    ("data", ["Australia", "12989", "https://data.csiro.au/collection/65918", "https://doi.org/10.25919/v2qx-vp27"]),
    ("data", ["TCMB", "20753", "https://tcmbank.cn/", "https://doi.org/10.1038/s41392-023-01339-1"]),
    ("data", ["CMNPD", "21387", "https://www.cmnpd.org/browse/compound", "https://doi.org/10.1093/nar/gkaa763"]),
    ("data", ["ATLAS", "27200", "https://www.npatlas.org/about", "https://doi.org/10.1021/acscentsci.9b00806"]),
    ("data", ["CMAUP", "41190", "https://bidd.group/CMAUP/index.html", "https://doi.org/10.1093/nar/gkad921"]),
    ("data", ["IBIS", "64723", "https://www.ibscreen.com/natural-compounds", "https://doi.org/10.1080/07391102.2023.2175377"]),
    ("data", ["ZINC15_BIO", "225676", "https://zinc15.docking.org/substances/subsets/biogenic/", "https://pubs.acs.org/doi/abs/10.1021/acs.jcim.5b00559"]),
    ("section", "Repositioning Database"),
    ("data", ["ZINC15_W", "5007", "https://zinc15.docking.org/substances/subsets/world/", "https://pubs.acs.org/doi/abs/10.1021/acs.jcim.5b00559"]),
]


def _add_step1_section(document: Any, job_dir: str, state: dict[str, Any], idioma: str = "pt") -> bool:
    step1 = state.get("step1")
    if not isinstance(step1, dict) or not step1:
        return False

    lang = "en" if idioma == "en" else "pt"
    texts = _STEP1_TEXTS

    target_chembl_id = step1.get("target_chembl_id", "")
    organism = step1.get("organism_name", "")
    target_type = step1.get("target_type", "")
    assay_type = step1.get("assay_type", "")
    metrics = step1.get("assay_metric_selected", [])
    units = step1.get("assay_unit_selected", [])
    included = step1.get("assay_description_included", "")
    excluded = step1.get("assay_description_excluded", "")

    internal_dir = os.path.join(job_dir, "DATA_BASES", "INTERNAL_DATA")
    initial_path = _latest_glob(os.path.join(internal_dir, f"df1_by_activity_{target_chembl_id}_{organism}*.csv"))
    base_path = _latest_glob(os.path.join(internal_dir, f"df1_base_{target_chembl_id}_{organism}*.csv"))
    initial_count = _row_count(initial_path)
    base_count = _row_count(base_path)

    _bar(document, "STEP 1 - DATASET PREPARATION", COLOR_SECTION_BAR)

    # Fixed intro paragraph (internal vs external datasets) - same wording for every job.
    _para(document).add_run(texts["intro_pt"] if lang == "pt" else texts["intro_en"])

    p = _para(document)

    def add(text: str, bold: bool = False) -> None:
        _add(p, text, bold)

    def add_bi(text_pt: str, text_en: str, bold: bool = False) -> None:
        add(text_pt if lang == "pt" else text_en, bold)

    add_bi(
        "Para compor o dataset interno, foi usada a base de dados ",
        "The internal dataset used to screen the supervised machine learning models was built from the ",
    )
    add("ChEMBL", True)
    add_bi(
        ", por meio da biblioteca python chembl_webresource_client. O Target Type selecionado foi ",
        " database via the chembl_webresource_client Python library. The Target Type selected was ",
    )
    add(target_type or "n/a", True)
    add_bi(" e, para Organism name, empregado o termo ", " and, for Organism name, the term ")
    add(organism or "n/a", True)
    add_bi(". Utilizando o Target ChEMBL ID ", ". Using Target ChEMBL ID ")
    add(target_chembl_id or "n/a", True)
    if initial_count is not None:
        add_bi(
            " obteve-se um dataframe (df1_by_activity) contendo ",
            ", a dataframe was obtained (df1_by_activity) containing ",
        )
        add(f"{initial_count:,}", True)
        add_bi(
            " compostos com dados experimentais de bioatividade. ",
            " compounds with experimental bioactivity data. ",
        )
    else:
        add(". ", False)
    add_bi(
        "Este dataframe inicial passou por filtragens onde foram selecionados Assay Type ",
        "This initial dataframe was filtered by Assay Type ",
    )
    add(assay_type or "n/a", True)
    if metrics:
        add(", Assay Metric ", False)
        _join_bold_list_bi(add, metrics, lang)
    if units:
        add(", Assay Unit ", False)
        _join_bold_list_bi(add, units, lang)
    add(". ", False)
    if included or excluded:
        add_bi(
            "Assay description foi usado para a curadoria criteriosa das informações dos bioensaios, "
            "de forma a evitar discrepâncias significativas na natureza destes que pudessem "
            "comprometer o estabelecimento de correlações entre os descritores e os valores de "
            "bioatividade. Esse processo permitiu a seleção dos termos de inclusão ",
            "Assay description was used for careful curation of the bioassay information, in order "
            "to avoid significant discrepancies in their nature that could compromise the "
            "establishment of correlations between the descriptors and the bioactivity values. "
            "This process allowed the selection of the inclusion terms ",
        )
        if included:
            add(included, True)
        if excluded:
            add_bi(" e dos termos de exclusão ", ", and the exclusion terms ")
            add(excluded, True)
        add(". ", False)
    else:
        add_bi(
            "Nenhum termo de filtragem, por inclusão ou exclusão, foi empregado em assay "
            "description, pois implicaria em uma redução significativa dos dados, que já nessa "
            "etapa inicial podem ser considerados escassos. ",
            "No assay description filtering term (inclusion or exclusion) was applied, since it "
            "would imply a significant reduction in the data, which at this initial stage may "
            "already be considered scarce. ",
        )
    if base_count is not None:
        add_bi(
            "Ao final dessa etapa o dataframe original foi reduzido para ",
            "At the end of this stage the original dataframe was reduced to ",
        )
        add(f"{base_count:,}", True)
        add_bi(" compostos (df1_base).", " compounds (df1_base).")

    # Fixed external-dataset-criteria paragraph (references Table 1) - same wording for every job.
    _para(document).add_run(texts["external_criteria_pt"] if lang == "pt" else texts["external_criteria_en"])
    document.add_paragraph()

    _field_line(document, texts["cell_line"][lang], step1.get("cell_name", ""))
    _field_line(document, texts["cell_chembl_id"][lang], step1.get("cell_chembl_id", ""))

    _add_caption(document, texts["table1_caption_pt"] if lang == "pt" else texts["table1_caption_en"])
    _add_static_table(document, _TABLE_1_HEADER, _TABLE_1_ROWS, col_ratios=[0.18, 0.14, 0.34, 0.34])
    footnote = document.add_paragraph()
    footnote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footnote_run = footnote.add_run(texts["table1_footnote_pt"] if lang == "pt" else texts["table1_footnote_en"])
    footnote_run.font.size = Pt(9)

    document.add_page_break()
    return True


# --------------------------------------------------------------------------------------
# Section 2: STEP 2 - Preprocessing and Exploratory Analysis (includes Generating
# Categories/Druggability Descriptors, moved here from the former STEP 3)
# --------------------------------------------------------------------------------------

_STEP2_TEXTS = {
    "descriptive_boilerplate_pt": (
        "Estatísticas descritivas, cálculo do coeficiente de assimetria de Fisher-Pearson, "
        "cálculo da curtose e diferentes representações gráficas (Histplot, Boxplot, Q-Q plot e "
        "Violin plot) também foram utilizadas a fim de verificarmos assimetrias na distribuição "
        "dos dados e possível presença de outliers."
    ),
    "descriptive_boilerplate_en": (
        "Descriptive statistics, Fisher-Pearson skewness coefficient calculation, kurtosis "
        "calculation, and different graphical representations (Histplot, Boxplot, Q-Q plot, and "
        "Violin plot) were also used in order to verify asymmetries in the data distribution and "
        "possible presence of outliers."
    ),
    "figure2_caption_pt": "Figura 2. Análise exploratória das distribuições através de gráficos de histograma (Acima) e boxplot (Abaixo), comparando os valores não transformados com os valores transformados.",
    "figure2_caption_en": "Figure 2. Exploratory analysis of the distributions through histogram (Above) and boxplot (Below) charts, comparing the non-transformed values with the transformed values.",
    "table2_caption_pt": "Tabela 2. Estatística descritiva para os valores de {col}.",
    "table2_caption_en": "Table 2. Descriptive statistics for the {col} values.",
    "figure3_caption_pt": "Figura 3. Número de compostos por classe. Mostrando o balanceamento entre as classes que poderão ser usadas em modelos de classificação.",
    "figure3_caption_en": "Figure 3. Number of compounds per class. Shows the balance between the classes that may be used in classification models.",
}

_TRANSFORM_LABELS = {
    "-Log10 (Only > 0 - pIC50|pMIC)": {"pt": "logaritmo inverso (-log)", "en": "inverse logarithm (-log)"},
    "ln (Only > 0)": {"pt": "logaritmo natural (ln)", "en": "natural logarithm (ln)"},
    "Square root (Only >= 0)": {"pt": "raiz quadrada", "en": "square root"},
    "Cubic root (All values)": {"pt": "raiz cúbica", "en": "cubic root"},
    "Box-Cox (Only > 0)": {"pt": "Box-Cox", "en": "Box-Cox"},
    "Yeo-Johnson (All values)": {"pt": "Yeo-Johnson", "en": "Yeo-Johnson"},
}

_OUTLIER_METHOD_LABELS = {
    "IQR": {"pt": "intervalo interquartílico (IQR)", "en": "interquartile range (IQR)"},
    "Z-Score": {"pt": "Z-Score", "en": "Z-Score"},
    "SD": {"pt": "desvio padrão (SD)", "en": "standard deviation (SD)"},
}

_OUTLIER_FORMULA_TEXTS = {
    "IQR": {
        "pt": "considerando-se outliers valores abaixo do limite inferior (Q1 - Threshold × IQR) ou acima do limite superior (Q3 + Threshold × IQR)",
        "en": "considering as outliers values below the lower bound (Q1 - Threshold × IQR) or above the upper bound (Q3 + Threshold × IQR)",
    },
    "Z-Score": {
        "pt": "considerando-se outliers valores com Z-Score abaixo de -Threshold ou acima de +Threshold",
        "en": "considering as outliers values with a Z-Score below -Threshold or above +Threshold",
    },
    "SD": {
        "pt": "considerando-se outliers valores abaixo da média menos Threshold desvios-padrão ou acima da média mais Threshold desvios-padrão",
        "en": "considering as outliers values below the mean minus Threshold standard deviations or above the mean plus Threshold standard deviations",
    },
}

_NORMAL_TEST_LABELS = {
    "normal_test_shapiro": "Shapiro-Wilk",
    "normal_test_anderson": "Anderson-Darling",
    "normal_test_kolmogorov": "Kolmogorov-Smirnov",
}


def _parse_outlier_report(text: str) -> Optional[dict[str, Any]]:
    """Parses the "{n} below,\\n {n} above\\n by interval [x, y]" (IQR/SD) or "... by threshold t"
    (Z-Score) report text written by CODRUG's run_outlier() into a dict of numbers."""
    match = re.search(
        r"(\d+)\s*below,\s*(\d+)\s*above\s*by\s*(?:interval\s*\[\s*([^,\]]+),\s*([^\]]+)\]|threshold\s+([\d.eE+-]+))",
        text,
    )
    if not match:
        return None
    n_below, n_above = int(match.group(1)), int(match.group(2))
    return {
        "n_below": n_below,
        "n_above": n_above,
        "n_total": n_below + n_above,
        "lower": match.group(3).strip() if match.group(3) else None,
        "upper": match.group(4).strip() if match.group(4) else None,
        "threshold_from_file": match.group(5),
    }


def _describe_col_stats(df) -> Optional[dict[str, str]]:
    """Reads a df_stats_<column>_*.csv (View Descriptive Statistics output) into a dict of
    formatted values, handling both the current column names (Q1/Q2/Q3) and the legacy ones
    (25%/50%/75%) written before that rename."""
    if df is None or df.empty:
        return None
    row = df.iloc[0]
    out = {}
    for key, legacy in (
        ("count", "count"), ("mean", "mean"), ("std", "std"), ("min", "min"),
        ("Q1", "25%"), ("Q2", "50%"), ("Q3", "75%"), ("max", "max"),
    ):
        if key in df.columns:
            out[key] = str(row[key])
        elif legacy in df.columns:
            out[key] = str(row[legacy])
    return out or None


def _add_step2_3_section(document: Any, job_dir: str, state: dict[str, Any], idioma: str = "pt") -> bool:
    step2 = state.get("step2") if isinstance(state.get("step2"), dict) else {}
    if not step2:
        return False

    lang = "en" if idioma == "en" else "pt"
    texts = _STEP2_TEXTS
    statistics_state = state.get("statistics") if isinstance(state.get("statistics"), dict) else {}

    stats_dir = os.path.join(job_dir, "RESULTS", "STATISTICS")
    internal_dir = os.path.join(job_dir, "DATA_BASES", "INTERNAL_DATA")

    _bar(document, "STEP 2 - Preprocessing and Exploratory Analysis", COLOR_SECTION_BAR)
    p = _para(document)

    def add(text: str, bold: bool = False) -> None:
        _add(p, text, bold)

    def add_bi(text_pt: str, text_en: str, bold: bool = False) -> None:
        add(text_pt if lang == "pt" else text_en, bold)

    # 1) Columns of interest kept - read directly from df2_ColumnFiltered's own header, since the
    # "1. Select columns of interest" list widget's saved selection reflects whatever dataframe
    # was loaded LAST in the session (not necessarily this filtering step), while the CSV it
    # wrote is always accurate.
    column_filtered_path = _latest_glob(os.path.join(internal_dir, "df2_ColumnFiltered_*.csv"))
    kept_columns = list(_read_csv(column_filtered_path, nrows=0).columns) if column_filtered_path else []
    if kept_columns:
        add_bi("No preprocessamento foram mantidas apenas as colunas de interesse ", "During preprocessing, only the columns of interest ", False)
        _join_bold_list_bi(add, kept_columns, lang)
        add_bi(". ", " were kept. ", False)

    # 2) Null/empty/out-of-range value removal.
    null_filtered_path = _latest_glob(os.path.join(internal_dir, "df2_NullEmptyFiltered_*.csv"))
    before_count = _row_count(column_filtered_path)
    after_null_count = _row_count(null_filtered_path)
    if before_count is not None and after_null_count is not None:
        removed_count = before_count - after_null_count
        add_bi("A seguir ", "Next, ", False)
        add(f"{removed_count:,}", True)
        add_bi(
            " compostos com value vazio, nulo ou fora do range para números em float64 foram eliminados do dataframe. ",
            " compounds with an empty, null, or out-of-range float64 value were removed from the dataframe. ",
            False,
        )

    # 3) Unit standardization.
    unit_col = step2.get("unit_column", "")
    unit_path = _latest_glob(os.path.join(internal_dir, "df2_unit_*.csv"))
    unit_count = _row_count(unit_path)
    if unit_col:
        add_bi("A unidade padrão foi definida como ", "The standard unit was defined as ", False)
        add(unit_col, True)
        add_bi(
            " e convertidas, quando possível, todas as unidades diferentes para este padrão. ",
            ", and all different units were converted to this standard whenever possible. ",
            False,
        )
        if unit_count is not None:
            add_bi("Ao final restaram ", "In the end, ", False)
            add(f"{unit_count:,}", True)
            add_bi(" compostos (df2_unit). ", " compounds remained (df2_unit). ", False)

    # 4) Deduplication.
    rep_method = step2.get("rep_method", "")
    rep_count = _row_count(_latest_glob(os.path.join(internal_dir, "df2_TratRepetitions_*.csv")))
    if rep_method:
        add_bi("Repetições do mesmo composto dentro do dataframe foram tratadas pelo método da ", "Repeated occurrences of the same compound within the dataframe were treated using the ", False)
        add(rep_method, True)
        add_bi(" mantendo uma única ocorrência", " method, keeping a single occurrence", False)
        if rep_count is not None:
            add_bi(", reduzindo o dataframe para ", ", reducing the dataframe to ", False)
            add(f"{rep_count:,}", True)
            add_bi(" compostos únicos (df2_TratRepetitions)", " unique compounds (df2_TratRepetitions)", False)
        add_bi(
            ", evitando-se assim que o mesmo composto, com os mesmos descritores, apresente valores de "
            "bioatividade muito discrepantes o que poderia dificultar a convergência dos modelos de ML com "
            "piora das métricas. ",
            ", thus avoiding the same compound, with the same descriptors, showing highly discrepant "
            "bioactivity values, which could hinder the convergence of the ML models and worsen their "
            "metrics. ",
            False,
        )

    # 5) Transformation.
    transform_col = step2.get("transform_column", "")
    transform_type = step2.get("transform_type", "")
    if transform_col and transform_type:
        transform_label = _TRANSFORM_LABELS.get(transform_type, {}).get(lang, transform_type)
        add_bi("O valor de ", "The ", False)
        add(transform_col, True)
        add_bi(" passou por transformação através do método de ", " value was transformed using the ", False)
        add(transform_label, True)
        if transform_type.startswith("-Log10"):
            add(f" (-log {transform_col})", False)
        add_bi(". ", " method. ", False)

    # 6) Normality tests + Figure 2 reference.
    normal_tests = [label for key, label in _NORMAL_TEST_LABELS.items() if statistics_state.get(key)]
    if normal_tests:
        add_bi("Foram aplicados os testes de ", "The ", False)
        _join_bold_list_bi(add, normal_tests, lang)
        add_bi(
            " para verificação da normalidade comparativa entre o valor original e transformado, essas "
            "distribuições podem ser observadas através da ",
            " test(s) were applied to check the comparative normality between the original and transformed "
            "values; these distributions can be observed in ",
            False,
        )
        add_bi("Figura 2", "Figure 2", True)
        add(". ", False)

    add_bi(texts["descriptive_boilerplate_pt"], texts["descriptive_boilerplate_en"], False)

    # Composite histogram/boxplot figure, if canonical PNGs are already on disk.
    metric_col = statistics_state.get("stat_column") or transform_col or unit_col
    candidates = []
    if metric_col:
        candidates += [
            os.path.join(stats_dir, f"histogram_{metric_col}.png"),
            os.path.join(stats_dir, f"histogram_{transform_col}.png") if transform_col else "",
            os.path.join(stats_dir, f"boxplot_{metric_col}.png"),
            os.path.join(stats_dir, f"boxplot_{transform_col}.png") if transform_col else "",
        ]
    composite_path = _composite_2x2_image([c for c in candidates if c], os.path.join(stats_dir, "_report_step2_distributions.png"))
    if composite_path:
        document.add_paragraph()
        _add_caption(document, texts["figure2_caption_pt"] if lang == "pt" else texts["figure2_caption_en"])
        _add_image(document, composite_path)

    # --- Outlier detection paragraph ---
    outlier_col = step2.get("outlier_column", "")
    outlier_method = step2.get("outlier_method", "")
    outlier_threshold = step2.get("outlier_threshold", "")
    if outlier_col and outlier_method:
        p2 = _para(document)

        def add2(text: str, bold: bool = False) -> None:
            _add(p2, text, bold)

        def add2_bi(text_pt: str, text_en: str, bold: bool = False) -> None:
            add2(text_pt if lang == "pt" else text_en, bold)

        method_label = _OUTLIER_METHOD_LABELS.get(outlier_method, {}).get(lang, outlier_method)
        formula_text = _OUTLIER_FORMULA_TEXTS.get(outlier_method, {}).get(lang, "")
        add2_bi("O método empregado para a detecção de outliers foi o ", "The method used for outlier detection was the ", False)
        add2(method_label, True)
        if formula_text:
            add2(f", {formula_text}", False)
        add2(". ", False)

        outlier_report_path = _latest_glob(os.path.join(stats_dir, f"{outlier_method}_Outliers_Removal_Report.txt"))
        parsed = None
        if outlier_report_path and os.path.isfile(outlier_report_path):
            try:
                parsed = _parse_outlier_report(Path(outlier_report_path).read_text(encoding="utf-8", errors="ignore"))
            except OSError:
                parsed = None

        stats_csv = _read_csv(_latest_glob(os.path.join(stats_dir, f"df_stats_{outlier_col}_*.csv")))
        col_stats = _describe_col_stats(stats_csv)

        add2_bi("Para ", "For ", False)
        add2(outlier_col, True)
        if outlier_threshold:
            add2_bi(" com Threshold de ", " with a Threshold of ", False)
            add2(str(outlier_threshold), True)
        if parsed:
            add2_bi(" um total de ", " a total of ", False)
            add2(f"{parsed['n_total']:,}", True)
            add2_bi(" outliers foram eliminados", " outliers were removed", False)
            add2_bi(", ", ", ", False)
            add2(str(parsed["n_below"]), True)
            add2_bi(" abaixo do limite inferior", " below the lower bound", False)
            if parsed.get("lower"):
                add2_bi(" de ", " of ", False)
                add2(parsed["lower"], True)
            add2_bi(" e ", " and ", False)
            add2(str(parsed["n_above"]), True)
            add2_bi(" acima do limite superior", " above the upper bound", False)
            if parsed.get("upper"):
                add2_bi(" de ", " of ", False)
                add2(parsed["upper"], True)
            if col_stats:
                add2_bi(" (Tabela 2)", " (Table 2)", False)
        add2(".", False)

        if col_stats:
            document.add_paragraph()
            caption_template = texts["table2_caption_pt"] if lang == "pt" else texts["table2_caption_en"]
            _add_caption(document, caption_template.format(col=outlier_col))
            stats_header = ["count", "mean", "std", "min", "Q1", "Q2", "Q3", "max"]
            stats_row = ("data", [col_stats.get(k, "") for k in stats_header])
            _add_static_table(document, stats_header, [stats_row])

    # --- Class categorization paragraph ---
    class_value_col = step2.get("class_value_column") or metric_col or "value"
    class_rules = []
    class1_name, class1_op, class1_ref = step2.get("class1_name"), step2.get("class1_operator"), step2.get("class1_reference")
    if class1_name and class1_op and class1_ref:
        class_rules.append((class1_name, f"{class_value_col} {class1_op} {class1_ref}"))
    class2_name, class2_min, class2_max = step2.get("class2_name"), step2.get("class2_min"), step2.get("class2_max")
    if class2_name and class2_min and class2_max:
        class_rules.append((class2_name, f"{class2_min} <= {class_value_col} <= {class2_max}"))
    class3_name, class3_op, class3_ref = step2.get("class3_name"), step2.get("class3_operator"), step2.get("class3_reference")
    if class3_name and class3_op and class3_ref:
        class_rules.append((class3_name, f"{class_value_col} {class3_op} {class3_ref}"))
    class_freq_path = None
    if class_rules:
        p3 = _para(document)

        def add3(text: str, bold: bool = False) -> None:
            _add(p3, text, bold)

        def add3_bi(text_pt: str, text_en: str, bold: bool = False) -> None:
            add3(text_pt if lang == "pt" else text_en, bold)

        add3_bi(
            "Visando permitir futuras abordagens por modelos de classificação, e análises de correlação entre "
            "classes e diferentes variáveis, foi criada uma coluna ",
            "Aiming to enable future classification-model approaches and correlation analyses between classes "
            "and different variables, a ",
            False,
        )
        add3("Class", True)
        add3_bi(", categorizando os compostos como ", " column was created, categorizing compounds as ", False)
        _join_bold_list_bi(add3, [name for name, _ in class_rules], lang)
        add3_bi(" (", " (", False)
        for index, (name, rule) in enumerate(class_rules):
            if index > 0:
                add3("; ", False)
            add3(f"{name}: {rule}", True)
        add3_bi(
            "); compostos que não corresponderam a nenhuma dessas regras ficaram sem classificação (",
            "); compounds matching none of these rules were left unclassified (",
            False,
        )
        add3_bi("Figura 3", "Figure 3", True)
        add3(").", False)

        class_freq_path = _latest_glob(os.path.join(job_dir, "RESULTS", f"freq_{class_value_col}*.png"))
        if class_freq_path:
            document.add_paragraph()
            _add_caption(document, texts["figure3_caption_pt"] if lang == "pt" else texts["figure3_caption_en"])
            _add_image(document, class_freq_path)

    # --- Druggability descriptors paragraph ---
    druggability_pairs = [
        ("druggability_mw", "MW", "druggability_mw_min", "druggability_mw_max"),
        ("druggability_logp", "LogP", "druggability_logp_min", "druggability_logp_max"),
        ("druggability_hdonor", "H-Donors", "druggability_hdonor_min", "druggability_hdonor_max"),
        ("druggability_haceptor", "H-Acceptors", "druggability_haceptor_min", "druggability_haceptor_max"),
        ("druggability_tpsa", "TPSA", "druggability_tpsa_min", "druggability_tpsa_max"),
        ("druggability_rbonds", "Rotatable Bonds", "druggability_rbonds_min", "druggability_rbonds_max"),
        ("druggability_ro5", "RO5 Violations", "druggability_ro5_min", "druggability_ro5_max"),
    ]
    active_ranges = [
        f"{label} ({step2.get(min_k)} to {step2.get(max_k)})"
        for flag_k, label, min_k, max_k in druggability_pairs
        if step2.get(flag_k)
    ]
    if active_ranges:
        p4 = _para(document)

        def add4(text: str, bold: bool = False) -> None:
            _add(p4, text, bold)

        def add4_bi(text_pt: str, text_en: str, bold: bool = False) -> None:
            add4(text_pt if lang == "pt" else text_en, bold)

        add4_bi("Foram calculados os descritores de drogabilidade ", "Druggability descriptors were computed for ", False)
        add4(", ".join(active_ranges), True)
        add4_bi(
            " para análise exploratória e aplicação de estatísticas comparativas entre classes e correlação de "
            "variáveis no dataframe, porém os dados não foram filtrados por drogabilidade nesta etapa.",
            " for exploratory analysis and comparative statistics between classes and variable correlation in "
            "the dataframe, but the data was not filtered by druggability at this stage.",
            False,
        )

    document.add_paragraph()
    return True


# --------------------------------------------------------------------------------------
# Section 3: STEP 3 - Features Engineering (internal name/JSON key "step4" predates the STEP3
# removal/renumbering - see feedback_codrug_refactor_scope convention)
# --------------------------------------------------------------------------------------

_STEP3_TEXTS = {
    "descriptors_intro_pt": (
        "Na etapa de engenharia de features foram gerados descritores a partir das informações "
        "estruturais presentes na coluna canonical_smiles e assim obtidas as variáveis "
        "independentes (valores de x ou features), usadas na construção dos modelos de ML."
    ),
    "descriptors_intro_en": (
        "In the feature engineering stage, descriptors were generated from the structural "
        "information in the canonical_smiles column, obtaining the independent variables (x "
        "values, or features) used to build the ML models."
    ),
}

_DESCRIPTOR_CATEGORY_LABELS = {
    "descriptors_1d2d": {"pt": "1D/2D", "en": "1D/2D"},
    "descriptors_3d": {"pt": "3D", "en": "3D"},
    "descriptors_fingerprint": {"pt": "Fingerprints", "en": "Fingerprints"},
}

_PREP_FLAG_LABELS = {
    "prep_remove_salt": {"pt": "remoção de sais", "en": "salt removal"},
    "prep_aromaticity": {"pt": "detecção de aromaticidade", "en": "aromaticity detection"},
    "prep_tautomers": {"pt": "padronização de tautômeros", "en": "tautomer standardization"},
    "prep_nitro": {"pt": "padronização de grupos nitro", "en": "nitro group standardization"},
}

# Mirrors CODRUG.py's SELECTION_METHOD_ROWS/PROJ_METHOD_ROWS (STEP 4's shared "Parameters:"
# table): which row(s) of state["step4"]["projection_parameters_table"] are actually relevant
# for each Selection/Projection method. Kept as a parallel list here (report generator has no
# PyQt/CODRUG.py import) rather than shared code - if those dicts change, update both.
_SELECTION_METHOD_PARAMS = {
    "Variance Threshold": ["Variance threshold"],
    "Correlation Threshold": ["corr_threshold"],
    "SelectKBest (ANOVA F-test)": ["top_k"],
    "SelectKBest (Chi2)": ["top_k"],
    "SelectPercentile": [],
    "Mutual Information": ["top_k"],
    "RFE": ["top_k"],
    "SFS (Forward)": ["top_k"],
    "SFS (Backward)": ["top_k"],
    "Lasso (L1)": ["top_k", "alpha"],
    "Ridge (L2)": ["top_k", "alpha"],
    "Elastic Net": ["top_k", "alpha"],
    "Tree-based Importance (RandomForest)": ["top_k"],
    "GB Importance": ["top_k"],
    "Stepwise AIC": [],
    "Stepwise BIC": [],
}
_PROJECTION_METHOD_PARAMS = {
    "PCA": ["Number of Components"],
    "Kernel PCA": ["Number of Components", "kernel", "gamma"],
    "t-SNE": ["Number of Components", "perplexity", "learning_rate", "n_iter"],
    "UMAP": ["Number of Components", "n_neighbors", "min_dist"],
    "TruncatedSVD": ["Number of Components"],
    "NMF": ["Number of Components"],
    "Isomap": ["Number of Components", "n_neighbors"],
    "Spectral Embedding": ["Number of Components", "affinity"],
    "LDA": ["Number of Components"],
    "PLS Regression": ["Number of Components"],
    "PLS-DA": ["Number of Components"],
}


def _col_count(path: Optional[str]) -> Optional[int]:
    df = _read_csv(path, nrows=0)
    return None if df is None else len(df.columns)


def _param_label(key: str) -> str:
    """Number of Components / Variance threshold are already human-readable; the rest
    (top_k, corr_threshold, alpha, kernel, gamma, ...) are raw table-row keys."""
    if key in ("Number of Components", "Variance threshold"):
        return key
    return key.replace("_", " ").title()


def _add_step4_section(document: Any, job_dir: str, state: dict[str, Any], idioma: str = "pt") -> bool:
    step4 = state.get("step4")
    if not isinstance(step4, dict) or not step4:
        return False

    scaling_method = step4.get("scaling_method") or step4.get("last_scaling_method")
    selection_method = step4.get("last_selection_method")
    projection_method = step4.get("last_projection_method")
    descriptors = step4.get("descriptors_selected")
    has_descriptors = isinstance(descriptors, dict) and bool(descriptors.get("selected"))
    if not (has_descriptors or scaling_method or selection_method or projection_method):
        return False

    lang = "en" if idioma == "en" else "pt"
    texts = _STEP3_TEXTS
    internal_dir = os.path.join(job_dir, "DATA_BASES", "INTERNAL_DATA")
    param_table = step4.get("projection_parameters_table") if isinstance(step4.get("projection_parameters_table"), dict) else {}

    _bar(document, "STEP 3 - Features Engineering", COLOR_SECTION_BAR)
    p = _para(document)

    def add(text: str, bold: bool = False) -> None:
        _add(p, text, bold)

    def add_bi(text_pt: str, text_en: str, bold: bool = False) -> None:
        add(text_pt if lang == "pt" else text_en, bold)

    add_bi(texts["descriptors_intro_pt"], texts["descriptors_intro_en"], False)

    # Descriptor generation: category (1D/2D, 3D, Fingerprints) + specific descriptor name(s),
    # PaDELPy version, and the pre-processing flags actually enabled.
    descriptors_path = _latest_glob(os.path.join(internal_dir, "df3_descriptors_*.csv")) or \
        _latest_glob(os.path.join(internal_dir, "df4_descriptors_*.csv"))
    if has_descriptors:
        categories = [
            _DESCRIPTOR_CATEGORY_LABELS[key][lang]
            for key in ("descriptors_1d2d", "descriptors_3d", "descriptors_fingerprint")
            if step4.get(key)
        ]
        add_bi(" O descritor selecionado foi do tipo ", " The selected descriptor type was ", False)
        _join_bold_list_bi(add, categories, lang)
        add(", ", False)
        _join_bold_list_bi(add, descriptors["selected"], lang)
        add_bi(", gerado através do Padelpy", ", generated via PaDELPy", False)
        try:
            import importlib.metadata as _importlib_metadata
            padelpy_version = _importlib_metadata.version("padelpy")
        except Exception:
            padelpy_version = None
        if padelpy_version:
            add_bi(" (PaDEL-Descriptor versão ", " (PaDEL-Descriptor version ", False)
            add(padelpy_version, True)
            add(")", False)
        prep_labels = [_PREP_FLAG_LABELS[key][lang] for key in _PREP_FLAG_LABELS if step4.get(key)]
        if prep_labels:
            add_bi(" com ", " with ", False)
            _join_bold_list_bi(add, prep_labels, lang)
        add(". ", False)

        if descriptors_path:
            desc_df = _read_csv(descriptors_path, nrows=0)
            if desc_df is not None:
                exclude = {
                    step4.get("structure_column"), step4.get("name_column"),
                    step4.get("bioactivity_column"), step4.get("class_column"),
                }
                descriptor_cols = [c for c in desc_df.columns if c not in exclude]
                if descriptor_cols:
                    add_bi(
                        "Este descritor fornece códigos binários que identificam a presença (1) ou "
                        "ausência (0) de ",
                        "This descriptor provides binary codes identifying the presence (1) or "
                        "absence (0) of ",
                        False,
                    )
                    add(f"{len(descriptor_cols):,}", True)
                    add_bi(
                        " subestruturas dentre os compostos do nosso dataframe (",
                        " substructures among the compounds in our dataframe (",
                        False,
                    )
                    add(f"{descriptor_cols[0]} {'a' if lang == 'pt' else 'to'} {descriptor_cols[-1]}", True)
                    add("). ", False)

    # Scaling/Selection/Projection are only reported if their respective output file actually
    # exists in THIS job's INTERNAL_DATA - state["step4"]'s scaling_method/selection_method/
    # projection_method can be stale (e.g. carried over from a different job that was never
    # cleared), so the method name alone isn't proof the step was actually run for this job.
    scaling_path = _latest_glob(os.path.join(internal_dir, "df3_scaling_*.csv")) or \
        _latest_glob(os.path.join(internal_dir, "df4_scaling_*.csv"))
    selection_path = _latest_glob(os.path.join(internal_dir, "df3_selection_*.csv")) or \
        _latest_glob(os.path.join(internal_dir, "df4_selection_*.csv"))
    projection_path = _latest_glob(os.path.join(internal_dir, "df3_projection_*.csv")) or \
        _latest_glob(os.path.join(internal_dir, "df4_projection_*.csv"))

    if scaling_method and scaling_path:
        add_bi("Para escalonamento das features foi empregado o método ", "Feature scaling was performed using the ", False)
        add(scaling_method, True)
        add_bi(". ", " method. ", False)

    if selection_method and selection_path:
        n_before = _col_count(descriptors_path)
        n_after = _col_count(selection_path)
        add_bi(
            "Para redução de dimensionalidade, com vistas a redução do risco de multicolinearidade "
            "e Overfitting nos modelos, na seleção de features foi aplicado o método de ",
            "For dimensionality reduction, aiming to reduce the risk of multicollinearity and "
            "Overfitting in the models, the ",
            False,
        )
        add(selection_method, True)
        add_bi("", " method was applied for feature selection", False)
        if n_before is not None and n_after is not None:
            add_bi(", o que reduziu de ", ", reducing ", False)
            add(f"{n_before:,}", True)
            add_bi(" para ", " to ", False)
            add(f"{n_after:,}", True)
            add_bi(" colunas de descritores", " descriptor columns", False)
        add(". ", False)
        for param_key in _SELECTION_METHOD_PARAMS.get(selection_method, []):
            value = param_table.get(param_key)
            if value:
                add_bi(f"O parâmetro {_param_label(param_key)} foi predefinido em ", f"The {_param_label(param_key)} parameter was set to ", False)
                add(str(value), True)
                add(". ", False)

    if projection_method and projection_path:
        n_comp = param_table.get("Number of Components")
        add_bi("Já para a projeção de variáveis latentes o método de ", "For latent variable projection, the ", False)
        add(projection_method, True)
        add_bi(" foi aplicado", " method was applied", False)
        if n_comp:
            add_bi(" com obtenção de ", " obtaining ", False)
            add(str(n_comp), True)
            add_bi(" componentes", " components", False)
        add(". ", False)
        for param_key in _PROJECTION_METHOD_PARAMS.get(projection_method, []):
            if param_key == "Number of Components":
                continue
            value = param_table.get(param_key)
            if value:
                add_bi(f"O parâmetro {_param_label(param_key)} foi predefinido em ", f"The {_param_label(param_key)} parameter was set to ", False)
                add(str(value), True)
                add(". ", False)

    # External dataset column-count note.
    ext_dir = os.path.join(job_dir, "DATA_BASES", "EXTERNAL_DATA")
    reference_path = selection_path or descriptors_path
    if reference_path:
        reference_df = _read_csv(reference_path, nrows=1)
        for ext_path in sorted(glob.glob(os.path.join(ext_dir, "df_external_descriptors_*.csv"))):
            ext_df = _read_csv(ext_path, nrows=1)
            if ext_df is None or reference_df is None:
                continue
            diff = len(set(ext_df.columns).symmetric_difference(set(reference_df.columns)))
            add_bi(
                f" O dataset externo {os.path.basename(ext_path)} apresentou {diff} colunas "
                f"diferentes em relação ao dataset interno {os.path.basename(reference_path)}, o "
                "que já era esperado, e essas colunas foram ignoradas no momento da aplicação dos "
                "melhores modelos para predição da bioatividade.",
                f" The external dataset {os.path.basename(ext_path)} presented {diff} different "
                f"columns relative to the internal dataset {os.path.basename(reference_path)}, as "
                "expected; these columns were ignored when applying the best models for bioactivity "
                "prediction.",
                False,
            )

    document.add_paragraph()
    return True


# --------------------------------------------------------------------------------------
# Section 4: STEP 4 - Machine Learning Models (Scikit-learn) (internal name/JSON key
# "step6_sklearn" predates the STEP3 removal/renumbering)
# --------------------------------------------------------------------------------------

_STEP4_ML_TEXTS = {
    "screening_intro_pt": "A triagem dos modelos de ML foi realizada com a biblioteca Scikit-learn. Foram avaliados {count} modelos: ",
    "screening_intro_en": "Model screening (Scikit-learn) was performed with the library. A total of {count} models were evaluated: ",
    "ranked_by_pt": " A métrica empregada para ranquear os modelos foi ",
    "ranked_by_en": " Models were ranked by ",
    "test_train_split_pt": " O conjunto teste representa {test}% e o conjunto treino {train}% do conjunto total de dados. ",
    "test_train_split_en": " The test set represented {test}% and the training set {train}% of the total data. ",
    "best_model_pt": "O modelo de melhor performance foi ",
    "best_model_en": "The best performing model was ",
    "tuning_pt": "Foi empregado o Hyperparameter Tuning sobre o modelo {model} usando {method} e {folds} Folds. Os hiperparâmetros e seus valores ótimos podem ser visualizados na tabela a seguir:",
    "tuning_en": "Hyperparameter Tuning was applied to the {model} model using {method} with {folds} Folds. The hyperparameters and their optimal values can be seen in the table below:",
    "cv_pt": "A validação cruzada do modelo {model} foi realizada através do método {method} com {folds} Folds, gerando a tabela a seguir:",
    "cv_en": "Cross-validation of the {model} model was performed using the {method} method with {folds} Folds, generating the table below:",
    "charts_pt": "Os gráficos gerados para o modelo {model} são apresentados a seguir:",
    "charts_en": "The charts generated for the {model} model are presented below:",
    "figure4_caption_pt": "Figura 4. Divisão dos conjuntos de treino e teste para o dataframe {df}.",
    "figure4_caption_en": "Figure 4. Train and test set split for the {df} dataframe.",
    "table3_caption_pt": "Tabela 3. Resultado da triagem dos modelos de machine learning para o dataframe interno {df}.",
    "table3_caption_en": "Table 3. Machine learning model screening results for the {df} internal dataframe.",
}

# Same shading used in the methodology template's STEP 4 tables (screening/tuning/cross-
# validation): header = COLOR_TITLE_BAR (already used for the top-level section bars too),
# highlighted first column = light blue, Cross-Validation's "OVERALL" summary row = yellow.
_STEP4_COLOR_HIGHLIGHT = "DEEBF6"
_STEP4_COLOR_OVERALL = "FEF2CC"

_STEP4_CHART_KINDS = ["predicted_vs_actual", "residuals", "manifold", "feature_importance"]


def _find_skl_csv(data_dir: str, prefix: str, usi: str, model_name: Optional[str] = None) -> Optional[str]:
    """Strict filename match: {prefix}_[{model_name}_]{usi}_{8 digits}_{6 digits}.csv - anchored
    (rather than a loose glob) so a model name that's a prefix of another one
    (KNeighborsRegressor / KNeighborsRegressor_1) can't cross-match the wrong file. Picks the
    most recently modified match."""
    if not os.path.isdir(data_dir):
        return None
    middle = f"{re.escape(model_name)}_" if model_name else ""
    pattern = re.compile(rf"^{re.escape(prefix)}_{middle}{re.escape(usi)}_\d{{8}}_\d{{6}}\.csv$")
    matches = [os.path.join(data_dir, f) for f in os.listdir(data_dir) if pattern.match(f)]
    return max(matches, key=os.path.getmtime) if matches else None


def _find_skl_image(midia_dir: str, model_name: str, kind: str, usi: str) -> Optional[str]:
    """Model result chart: <model>_<kind>_<usi>.png (current naming convention) - falls back to
    the older <model>_<kind>.png (no USI suffix) for charts saved before that convention was
    added, so existing jobs' images are still found."""
    for candidate in (
        os.path.join(midia_dir, f"{model_name}_{kind}_{usi}.png"),
        os.path.join(midia_dir, f"{model_name}_{kind}.png"),
    ):
        if os.path.isfile(candidate):
            return candidate
    return None


def _discover_skl_tuned_models(data_dir: str, usi: str) -> list[str]:
    """Every model that has a Hyperparameter Tuning results CSV for this USI, discovered
    directly from skl_tune_<model>_<usi>_<8 digits>_<6 digits>.csv filenames - rather than
    trusting tuning_settings' keys, which don't always match the model name actually baked into
    the file (e.g. a tuning_settings key "LGBMRegressor" whose file is really
    skl_tune_LGBMRegressor_1_<usi>_...csv once Tuning creates a numbered trained variant).
    Returns model names sorted alphabetically."""
    if not os.path.isdir(data_dir):
        return []
    pattern = re.compile(rf"^skl_tune_(.+)_{re.escape(usi)}_\d{{8}}_\d{{6}}\.csv$")
    models = {match.group(1) for f in os.listdir(data_dir) if (match := pattern.match(f))}
    return sorted(models)


def _lookup_skl_settings(settings: dict, model_name: str) -> dict:
    """settings[model_name] if present, else settings[<model_name without a trailing _N>] -
    tuning_settings/validation_settings are sometimes keyed by the base model name even though
    the actual result file (and this report's per-model loop) uses the numbered variant."""
    if model_name in settings:
        return settings[model_name]
    base = re.sub(r"_\d+$", "", model_name)
    return settings.get(base, {})


def _add_skl_results_table(
    document: Any, df, *,
    highlight_col: Optional[int] = None,
    highlight_col_bold: bool = False,
    overall_row_value: Optional[str] = None,
    max_rows: Optional[int] = None,
) -> None:
    """Screening/Tuning/Cross-Validation results table, styled to match the methodology
    template: header shaded COLOR_TITLE_BAR + bold; optionally one column (Model/Fold) shaded
    light blue (_STEP4_COLOR_HIGHLIGHT) throughout; the row whose first cell equals
    overall_row_value (e.g. "OVERALL") has its other cells shaded yellow
    (_STEP4_COLOR_OVERALL) and bold, matching the Cross-Validation table's summary row."""
    if df is None or df.empty:
        return
    cols = list(df.columns)
    rows = df if max_rows is None else df.head(max_rows)
    width = _content_width_cm(document)
    col_width = width / max(len(cols), 1)
    table = document.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col_idx, name in enumerate(cols):
        cell = table.cell(0, col_idx)
        _set_cell_width(cell, col_width)
        _shade_cell(cell, COLOR_TITLE_BAR)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].add_run(str(name)).bold = True
    for _, record in rows.iterrows():
        cells = table.add_row().cells
        is_overall = overall_row_value is not None and str(record[cols[0]]) == overall_row_value
        for col_idx, name in enumerate(cols):
            cell = cells[col_idx]
            _set_cell_width(cell, col_width)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            value = record[name]
            text = f"{value:.4g}" if isinstance(value, float) else str(value)
            run = cell.paragraphs[0].add_run(text)
            if highlight_col is not None and col_idx == highlight_col:
                _shade_cell(cell, _STEP4_COLOR_HIGHLIGHT)
                if highlight_col_bold or is_overall:
                    run.bold = True
            elif is_overall:
                _shade_cell(cell, _STEP4_COLOR_OVERALL)
                run.bold = True
    document.add_paragraph()


def _add_step5_section(document: Any, job_dir: str, state: dict[str, Any], idioma: str = "pt") -> bool:
    step6 = state.get("step6_sklearn") if isinstance(state.get("step6_sklearn"), dict) else {}
    if not step6:
        return False

    lang = "en" if idioma == "en" else "pt"
    texts = _STEP4_ML_TEXTS
    usi_root = os.path.join(job_dir, "RESULTS", "USI")
    if not os.path.isdir(usi_root):
        return False

    # One "STEP 4" repetition per available USI code (each is an independent Screening/Tuning/
    # Validation/Application run) - discovered directly from disk (RESULTS/USI/<code>/
    # skl_session.json), not from state["step6_sklearn"]["current_usi"] (a single "last selected"
    # pointer that can't represent more than one USI).
    usi_codes = sorted(
        name for name in os.listdir(usi_root)
        if os.path.isfile(os.path.join(usi_root, name, "skl_session.json"))
    )
    if not usi_codes:
        return False

    # Same descriptors dataframe referenced by STEP 3's own text (df3_descriptors_*.csv) - used
    # by the Figure 4/Table 3 captions below ("... for the <dataframe> dataframe").
    internal_dir = os.path.join(job_dir, "DATA_BASES", "INTERNAL_DATA")
    descriptors_path = _latest_glob(os.path.join(internal_dir, "df3_descriptors_*.csv")) or \
        _latest_glob(os.path.join(internal_dir, "df4_descriptors_*.csv"))
    descriptors_name = os.path.basename(descriptors_path) if descriptors_path else ""

    _bar(document, "STEP 4 - Machine Learning Models: Screening, Tuning, Validation and Application", COLOR_SECTION_BAR)

    added_any = False
    for usi in usi_codes:
        usi_dir = os.path.join(usi_root, usi)
        data_dir = os.path.join(usi_dir, "DATA")
        midia_dir = os.path.join(usi_dir, "MIDIA")
        try:
            session = json.loads(Path(os.path.join(usi_dir, "skl_session.json")).read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            continue

        usi_heading = document.add_paragraph()
        usi_heading.add_run(f"USI: {usi}").bold = True

        screening = session.get("screening_settings", {})
        best_model = session.get("best_model")

        # The models that actually took part in Screening are whatever the Screening results
        # CSV itself lists (its own "Model" column) - session["models"] instead reflects the
        # CURRENT set of saved/tuned model instances (grows over time as Tuning creates new
        # "_N" variants), which is a different, larger list and not what "participated in
        # screening" means.
        screening_csv = _find_skl_csv(data_dir, "skl_screening", usi)
        screening_df = _read_csv(screening_csv)
        screened_models = list(screening_df["Model"]) if screening_df is not None and "Model" in screening_df.columns else []

        p = _para(document)

        def add(text: str, bold: bool = False) -> None:
            _add(p, text, bold)

        def add_bi(text_pt: str, text_en: str, bold: bool = False) -> None:
            add(text_pt if lang == "pt" else text_en, bold)

        add_bi(
            texts["screening_intro_pt"].format(count=len(screened_models)),
            texts["screening_intro_en"].format(count=len(screened_models)),
            False,
        )
        if screened_models:
            _join_bold_list_bi(add, screened_models, lang)
        add(". ", False)
        if screening.get("sort_metric"):
            add_bi(texts["ranked_by_pt"], texts["ranked_by_en"], False)
            add(str(screening["sort_metric"]), True)
            add(". ", False)
        if screening.get("test_size"):
            test_size = float(screening["test_size"])
            add_bi(
                texts["test_train_split_pt"].format(test=f"{test_size * 100:.0f}", train=f"{100 - test_size * 100:.0f}"),
                texts["test_train_split_en"].format(test=f"{test_size * 100:.0f}", train=f"{100 - test_size * 100:.0f}"),
                False,
            )
        if best_model:
            add_bi(texts["best_model_pt"], texts["best_model_en"], False)
            add(str(best_model), True)
            add(". ", False)

        # Train/Test split chart ("View Frequency" button, Model Screening group), saved as
        # <Y column>_Train_Test_Plot_<USI>.png under this USI's own MIDIA folder (falls back to
        # the pre-USI-suffix name for older jobs).
        train_test_png = _latest_glob(os.path.join(midia_dir, f"*_Train_Test_Plot_{usi}.png")) or \
            _latest_glob(os.path.join(midia_dir, "*_Train_Test_Plot.png"))
        if train_test_png:
            _add_caption(document, texts["figure4_caption_pt" if lang == "pt" else "figure4_caption_en"].format(df=descriptors_name))
            _add_image(document, train_test_png, width_cm=8.0)

        if screening_df is not None:
            _add_caption(document, texts["table3_caption_pt" if lang == "pt" else "table3_caption_en"].format(df=descriptors_name))
            model_col = 0 if "Model" in screening_df.columns and screening_df.columns[0] == "Model" else None
            _add_skl_results_table(document, screening_df, highlight_col=model_col)

        tuning_settings = session.get("tuning_settings", {})
        # Discovered directly from skl_tune_<model>_<usi>_*.csv files on disk (not from
        # tuning_settings' own keys - see _discover_skl_tuned_models), so every model that went
        # through Tuning gets a table, including numbered variants like "BayesianRidge_1" whose
        # exact key isn't always present in tuning_settings.
        for model_name in _discover_skl_tuned_models(data_dir, usi):
            tune_csv = _find_skl_csv(data_dir, "skl_tune", usi, model_name)
            tune_df = _read_csv(tune_csv)
            if tune_df is None:
                continue
            if "rank_test_score" in tune_df.columns:
                # Only the best-ranked hyperparameter combination(s).
                tune_df = tune_df[tune_df["rank_test_score"] == 1]
            else:
                # No rank concept available (e.g. Optuna trial logs) - show just the first row.
                tune_df = tune_df.head(1)
            if tune_df.empty:
                continue
            settings = _lookup_skl_settings(tuning_settings, model_name)
            method = settings.get("method", "")
            folds = settings.get("folds", "")
            _para(document).add_run(
                texts["tuning_pt" if lang == "pt" else "tuning_en"].format(model=model_name, method=method, folds=folds)
            )
            _add_skl_results_table(document, tune_df)

        validation_settings = session.get("validation_settings", {})
        for model_name, settings in validation_settings.items():
            method = settings.get("method", "")
            folds = settings.get("folds", "")
            cv_csv = _find_skl_csv(data_dir, "skl_cv", usi, model_name)
            cv_df = _read_csv(cv_csv)
            if cv_df is not None:
                _para(document).add_run(
                    texts["cv_pt" if lang == "pt" else "cv_en"].format(model=model_name, method=method, folds=folds)
                )
                fold_col = 0 if list(cv_df.columns)[:1] == ["Fold"] else None
                _add_skl_results_table(document, cv_df, highlight_col=fold_col, highlight_col_bold=True, overall_row_value="OVERALL")

            chart_paths = [
                path for path in (_find_skl_image(midia_dir, model_name, kind, usi) for kind in _STEP4_CHART_KINDS)
                if path
            ]
            if chart_paths:
                _para(document).add_run(
                    texts["charts_pt" if lang == "pt" else "charts_en"].format(model=model_name)
                )
                for chart_path in chart_paths:
                    _add_image(document, chart_path)

        added_any = True
        document.add_paragraph()

    return added_any


# --------------------------------------------------------------------------------------
# Section 5: STEP 5 - Applicability Domain and Similarity Analysis (internal name/JSON key
# "step7_ad" predates the STEP3 removal/renumbering)
# --------------------------------------------------------------------------------------

_STEP5_TEXTS = {
    "leverage_pt": (
        "Para garantir a validade das predições dos modelos de ML treinados para o espaço "
        "químico em que foram utilizados, a análise do domínio de aplicabilidade foi realizada. "
        "O AD utiliza a técnica do intervalo de Hat (leverage) para identificar compostos fora "
        "do domínio de aplicabilidade. O leverage (h) de um composto é calculado a partir da "
        "matriz dos descritores utilizados no modelo. Um composto é considerado fora do domínio "
        "de aplicabilidade se seu leverage for maior que o limite crítico h*, destacado nos "
        "resultados por linhas pontilhadas verticais."
    ),
    "leverage_en": (
        "To ensure the validity of the predictions from the ML models trained for the chemical "
        "space in which they were used, the applicability domain analysis was performed. The AD "
        "uses the Hat (leverage) interval technique to identify compounds outside the "
        "applicability domain. The leverage (h) of a compound is calculated from the matrix of "
        "descriptors used in the model. A compound is considered outside the applicability "
        "domain if its leverage is greater than the critical limit h*, highlighted in the "
        "results by vertical dotted lines."
    ),
    "mahalanobis_pt": (
        "Adicionalmente, complementa-se a análise do domínio de aplicabilidade usando a "
        "distância de Mahalanobis e as propriedades estruturais dos compostos para a análise de "
        "similaridade. Este método considera tanto a similaridade química quanto a distância "
        "multidimensional entre os compostos no espaço dos descritores. Compostos identificados "
        "fora do domínio de aplicabilidade por qualquer um dos métodos foram filtrados do "
        "dataset externo para aplicação dos modelos nas previsões. Esta abordagem garantirá que "
        "o modelo QSAR desenvolvido seja robusto e aplicável dentro dos limites definidos."
    ),
    "mahalanobis_en": (
        "Additionally, the applicability domain analysis is complemented using the Mahalanobis "
        "distance and the compounds' structural properties for the similarity analysis. This "
        "method considers both the chemical similarity and the multidimensional distance "
        "between compounds in the descriptor space. Compounds identified as outside the "
        "applicability domain by either method were filtered from the external dataset before "
        "applying the models for the predictions. This approach will ensure the developed QSAR "
        "model is robust and applicable within the defined limits."
    ),
    "cutoffs_intro_pt": "A análise do domínio de aplicabilidade e da similaridade entre o dataframe interno e externo através dos métodos de ",
    "cutoffs_intro_en": "The applicability domain and similarity analysis between the internal and external dataframe through the ",
    "cutoffs_verdict_pt": " gerou o veredito de ",
    "cutoffs_verdict_en": " methods generated the verdict of ",
    "cutoffs_verdict_end_pt": " para cada composto do dataframe externo. ",
    "cutoffs_verdict_end_en": " for each compound in the external dataframe. ",
    "verdict_freq_intro_pt": "O dataframe ",
    "verdict_freq_intro_en": "The ",
    "verdict_freq_outro_pt": " apresentou a seguinte distribuição entre os tipos de veredito:",
    "verdict_freq_outro_en": " dataframe presented the following distribution among verdict types:",
    "composite_intro_pt": "Os gráficos de Williams, Mahalanobis, PCA_scatter e Tanimoto podem ser vistos a seguir:",
    "composite_intro_en": "The Williams, Mahalanobis, PCA_scatter and Tanimoto charts can be seen below:",
    "closing_pt": "Os compostos com veredito ",
    "closing_en": "Compounds with an ",
    "closing_mid_pt": " foram retirados dos dataframes externos antes da próxima etapa de análise consensual.",
    "closing_mid_en": " verdict were removed from the external dataframes before the next consensus analysis step.",
}

# Display order for verdict types (best to worst); only those actually present are shown.
_VERDICT_TYPE_ORDER = ["Within AD", "Borderline", "Outside AD"]


def _add_step7_section(document: Any, job_dir: str, state: dict[str, Any], idioma: str = "pt") -> bool:
    step7 = state.get("step7_ad")
    if not isinstance(step7, dict) or not step7:
        return False

    lang = "en" if idioma == "en" else "pt"
    texts = _STEP5_TEXTS
    midia_dir = os.path.join(job_dir, "RESULTS", "MIDIA")

    _bar(document, "STEP 5 - Applicability Domain and Similarity Analysis", COLOR_SECTION_BAR)

    _para(document).add_run(texts["leverage_pt"] if lang == "pt" else texts["leverage_en"])
    _para(document).add_run(texts["mahalanobis_pt"] if lang == "pt" else texts["mahalanobis_en"])

    p = _para(document)

    def add(text: str, bold: bool = False) -> None:
        _add(p, text, bold)

    def add_bi(text_pt: str, text_en: str, bold: bool = False) -> None:
        add(text_pt if lang == "pt" else text_en, bold)

    result_path = step7.get("result_csv")
    result_df = _read_csv(result_path)
    cutoffs = []
    if result_df is not None and not result_df.empty:
        cutoff_cols = {
            "leverage_cut": "leverage",
            "mahal_cut": "Mahalanobis",
            "knn_cut": "KNN",
            "tanimoto_cut": "Tanimoto",
        }
        for col, label in cutoff_cols.items():
            if col in result_df.columns:
                cutoffs.append((label, _fmt_num(result_df[col].iloc[0])))

    verdict_counts = step7.get("verdict_counts", {})
    verdict_types = [v for v in _VERDICT_TYPE_ORDER if v in verdict_counts] or list(verdict_counts.keys())

    if cutoffs and verdict_types:
        add_bi(texts["cutoffs_intro_pt"], texts["cutoffs_intro_en"], False)
        for index, (label, value) in enumerate(cutoffs):
            if index > 0:
                add(", " if index < len(cutoffs) - 1 else (" e " if lang == "pt" else " and "), False)
            add(f"{label} (cut-off = ", False)
            add(value, True)
            add(")", False)
        add_bi(texts["cutoffs_verdict_pt"], texts["cutoffs_verdict_en"], False)
        if len(verdict_types) > 1:
            add(", ".join(verdict_types[:-1]), False)
            add(" e " if lang == "pt" else " and ", False)
            add(verdict_types[-1], False)
        else:
            add(verdict_types[0], False)
        add_bi(texts["cutoffs_verdict_end_pt"], texts["cutoffs_verdict_end_en"], False)

    # Verdict distribution chart(s) ("Verdict Distribution" group, "View Frequency" button) -
    # one per external dataframe the user ran it for, named
    # Plot_verdict_freq_<external dataframe>_<timestamp>.png under RESULTS/MIDIA. Every match is
    # included (not just the latest), since the methodology template expects one distribution per
    # external dataframe (e.g. REPO and NP) rather than a single "current" one. The template
    # places the FIRST one right here (before the Williams/Mahalanobis/PCA/Tanimoto composite
    # figure) and any remaining ones after it.
    def _verdict_freq_entries():
        entries = []
        for verdict_freq_path in sorted(glob.glob(os.path.join(midia_dir, "Plot_verdict_freq_*.png"))):
            stem = os.path.splitext(os.path.basename(verdict_freq_path))[0]
            label_match = re.match(r"^Plot_verdict_freq_(.+)_\d{4}-\d{2}-\d{2}_\d{2}-\d{2}$", stem)
            ext_label = label_match.group(1) if label_match else stem[len("Plot_verdict_freq_"):]
            entries.append((ext_label, verdict_freq_path))
        return entries

    def _add_verdict_freq_entry(ext_label, path):
        entry_p = _para(document)
        add_pt, add_en = texts["verdict_freq_intro_pt"], texts["verdict_freq_intro_en"]
        _add(entry_p, add_pt if lang == "pt" else add_en, False)
        _add(entry_p, ext_label, True)
        _add(entry_p, texts["verdict_freq_outro_pt"] if lang == "pt" else texts["verdict_freq_outro_en"], False)
        _add_image(document, path)

    verdict_freq_entries = _verdict_freq_entries()
    if verdict_freq_entries:
        _add_verdict_freq_entry(*verdict_freq_entries[0])

    picks = []
    for prefix in ("Plot_williams_", "Plot_mahalanobis_", "Plot_pca_scatter_", "Plot_tanimoto_"):
        matches = sorted(glob.glob(os.path.join(midia_dir, f"{prefix}*.png")), key=os.path.getmtime, reverse=True)
        if matches:
            picks.append(matches[0])
    composite_path = _composite_2x2_image(picks, os.path.join(midia_dir, "_report_step7_ad.png"))
    if composite_path:
        _para(document).add_run(texts["composite_intro_pt"] if lang == "pt" else texts["composite_intro_en"])
        _add_image(document, composite_path)

    for ext_label, path in verdict_freq_entries[1:]:
        _add_verdict_freq_entry(ext_label, path)

    if verdict_counts:
        closing_p = _para(document)
        _add(closing_p, texts["closing_pt"] if lang == "pt" else texts["closing_en"], False)
        _add(closing_p, "Outside AD", True)
        _add(closing_p, texts["closing_mid_pt"] if lang == "pt" else texts["closing_mid_en"], False)

    document.add_paragraph()
    return True


# --------------------------------------------------------------------------------------
# Section 6: STEP 6 - Consensus Analysis (internal name/JSON key "step8_consensus" predates the
# STEP3 removal/renumbering)
# --------------------------------------------------------------------------------------

def _find_smiles_lookup(job_dir: str, ids: set[str]) -> dict[str, str]:
    """Best-effort ID -> SMILES lookup, scanning prediction CSVs under RESULTS/USI/**/PREDICTIONS
    (which carry Name/SMILES columns) for the small set of IDs the consensus hits table needs."""
    lookup: dict[str, str] = {}
    if not ids:
        return lookup
    pattern = os.path.join(job_dir, "RESULTS", "USI", "*", "PREDICTIONS", "*.csv")
    for path in glob.glob(pattern):
        if len(lookup) >= len(ids):
            break
        df = _read_csv(path)
        if df is None:
            continue
        id_col = next((c for c in ("Name", "name", "molecule_chembl_id", "ID", "id") if c in df.columns), None)
        smiles_col = next((c for c in ("SMILES", "smiles", "canonical_smiles") if c in df.columns), None)
        if not id_col or not smiles_col:
            continue
        for _, row in df[[id_col, smiles_col]].iterrows():
            key = str(row[id_col])
            if key in ids and key not in lookup:
                lookup[key] = str(row[smiles_col])
    return lookup


_STEP6_TEXTS = {
    "intro_pt": "A lista contendo o ranking dos compostos mais promissores, em ordem ",
    "intro_en": "The list containing the ranking of the most promising compounds, in ",
    "order_pt": {"decreasing": "decrescente", "increasing": "crescente"},
    "order_en": {"decreasing": "decreasing", "increasing": "increasing"},
    "of_values_pt": " dos valores de ",
    "of_values_en": " order of ",
    "reserved_pt": ", foi reservada para a análise consensual entre os modelos com melhores métricas selecionados pelo usuário: ",
    "reserved_en": " values, was reserved for the consensus analysis between the best-performing models selected by the user: ",
    "method_pt": ". Para a análise consensual entre os ranqueamentos gerados pelos melhores modelos foi utilizado o método de ",
    "method_en": ". The ",
    "method_mid_en": " method was used for the consensus analysis between the rankings generated by the best models",
    "cv_pt": ", e considerados apenas aqueles com coeficiente de variância entre os modelos menor ou igual a ",
    "cv_en": ", considering only those with a coefficient of variation between models less than or equal to ",
    "hits_pt": ". Foram obtidos ",
    "hits_en": ". A total of ",
    "hits_mid_pt": " Hits (",
    "hits_mid_en": " Hits (",
    "hits_end_pt": ") listados a seguir, com suas estruturas 2D:",
    "hits_end_en": ") were obtained, listed below with their 2D structures:",
    "table4_caption_pt": "Tabela 4. Resultado da análise consensual.",
    "table4_caption_en": "Table 4. Consensus analysis results.",
}


def _add_step8_section(document: Any, job_dir: str, state: dict[str, Any], idioma: str = "pt") -> bool:
    """Reads the "hits" CSV that STEP 6 itself already computed (df_consensus_hits_*.csv,
    referenced by step8_consensus.last_hits_file) rather than re-deriving the CV%/Hit% filter
    here - CODRUG.run_consensus_generate is the single source of truth for that logic, so the
    report can never drift out of sync with what the UI actually shows/saves."""
    step8 = state.get("step8_consensus")
    if not isinstance(step8, dict) or not step8:
        return False

    hits_path = step8.get("last_hits_file")
    if hits_path and os.path.isfile(hits_path):
        hits = _read_csv(hits_path)
    else:
        # Fall back to the full result only for older jobs saved before the hits CSV existed -
        # an existing hits file with 0 rows is a legitimate "no compound passed the filter"
        # outcome and must NOT fall back to the unfiltered table.
        hits = _read_csv(step8.get("last_result_file"))
    if hits is None or hits.empty:
        return False

    lang = "en" if idioma == "en" else "pt"
    texts = _STEP6_TEXTS
    method = step8.get("consensus_method", "Z-Score (Mean/SD)")
    cv_max = step8.get("cv_max_percent")
    hit_percent = step8.get("hit_percent")
    direction = "increasing" if step8.get("rank_increase") and not step8.get("rank_decrease") else "decreasing"

    # The models actually combined in the consensus analysis - derived from the slotN_value_column
    # fields (each one is "<Y column>_<model name>", e.g. "pIC50_LGBMRegressor_1"), stripping
    # their shared "<Y column>_" prefix to recover both the Y column name and the model list.
    slot_value_cols = [
        step8[k] for k in sorted(step8.keys())
        if re.match(r"^slot\d+_value_column$", k) and step8.get(k)
    ]
    y_col = ""
    selected_models = []
    if slot_value_cols:
        common_prefix = os.path.commonprefix(slot_value_cols)
        cut = common_prefix.rfind("_")
        if cut >= 0:
            y_col = common_prefix[:cut]
            selected_models = [v[cut + 1:] for v in slot_value_cols]
        else:
            selected_models = list(slot_value_cols)

    _bar(document, "STEP 6 - Consensus Analysis", COLOR_SECTION_BAR)
    p = _para(document)

    def add(text: str, bold: bool = False) -> None:
        _add(p, text, bold)

    def add_bi(text_pt: str, text_en: str, bold: bool = False) -> None:
        add(text_pt if lang == "pt" else text_en, bold)

    add_bi(texts["intro_pt"], texts["intro_en"], False)
    add(texts[f"order_{lang}"][direction], True)
    add_bi(texts["of_values_pt"], texts["of_values_en"], False)
    if y_col:
        add(y_col, True)
    add_bi(texts["reserved_pt"], texts["reserved_en"], False)
    if selected_models:
        _join_bold_list_bi(add, selected_models, lang)
    add_bi(texts["method_pt"], texts["method_en"], False)
    add(str(method), True)
    add_bi("", texts["method_mid_en"], False)
    if cv_max:
        add_bi(texts["cv_pt"], texts["cv_en"], False)
        add(f"{cv_max}%", True)
    add_bi(texts["hits_pt"], texts["hits_en"], False)
    add(f"{len(hits)}", True)
    add_bi(texts["hits_mid_pt"], texts["hits_mid_en"], False)
    if hit_percent:
        add(f"{hit_percent}%", True)
    add_bi(texts["hits_end_pt"], texts["hits_end_en"], False)
    document.add_paragraph()

    id_col = hits.columns[0]
    score_col = "consensus_score_mean" if "consensus_score_mean" in hits.columns else "zscore_consensus_mean"
    smiles_lookup = _find_smiles_lookup(job_dir, set(hits[id_col].astype(str)))

    _add_caption(document, texts["table4_caption_pt"] if lang == "pt" else texts["table4_caption_en"])

    width = _content_width_cm(document)
    col_widths = [width * 0.20, width * 0.30, width * 0.20, width * 0.15, width * 0.15]
    headers = ["Compound", "2D Structure", "Consensus Rank", "Consensus Score (mean)", "CV%"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col, (header, col_w) in enumerate(zip(headers, col_widths)):
        cell = table.cell(0, col)
        _set_cell_width(cell, col_w)
        _shade_cell(cell, COLOR_ACCENT_BAR)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].add_run(header).bold = True

    for _, record in hits.iterrows():
        compound_id = str(record[id_col])
        cells = table.add_row().cells
        _set_cell_width(cells[0], col_widths[0])
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[0].paragraphs[0].add_run(compound_id)
        _set_cell_width(cells[1], col_widths[1])
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        smiles = smiles_lookup.get(compound_id, "")
        image = molecule_image_bytes(smiles) if smiles else None
        if image is not None:
            cells[1].paragraphs[0].add_run().add_picture(image, width=Cm(col_widths[1] - 0.6))
        else:
            cells[1].paragraphs[0].add_run("N/A")
        _set_cell_width(cells[2], col_widths[2])
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[2].paragraphs[0].add_run(_fmt_num(record.get("consensus_rank", "")))
        _set_cell_width(cells[3], col_widths[3])
        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[3].paragraphs[0].add_run(_fmt_num(record.get(score_col, "")))
        _set_cell_width(cells[4], col_widths[4])
        cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[4].paragraphs[0].add_run(_fmt_num(record.get("bioactivity_cv_percent", "")))

    document.add_paragraph()
    return True


# --------------------------------------------------------------------------------------
# Report header ("RELATÓRIO FINAL CODRUG: Método e Resultados" topic, Job Name/Task Type/
# Started/Finished/Working Time) - like the Introduction, translated based on idioma.
# --------------------------------------------------------------------------------------

_METHOD_RESULTS_TEXTS = {
    "title": {
        "pt": "RELATÓRIO FINAL CODRUG: Método e Resultados",
        "en": "CODRUG FINAL REPORT: Method and Results",
    },
    "job_name": {"pt": "Nome do Projeto", "en": "Job Name"},
    "task_type": {"pt": "Tipo de Tarefa", "en": "Task Type"},
    "started": {"pt": "Iniciado", "en": "Started"},
    "finished": {"pt": "Finalizado", "en": "Finished"},
    "working_time": {"pt": "Tempo de Trabalho", "en": "Working Time"},
    "days": {"pt": "Dias", "en": "Days"},
    "hours": {"pt": "Horas", "en": "Hours"},
    "minutes": {"pt": "Minutos", "en": "Minutes"},
    "seconds": {"pt": "Segundos", "en": "Seconds"},
}

# CONFIG's Task Type value ("regression"/"classification"/"clustering") - only translated for
# display in the report header; the JSON always keeps the original English value.
_TASK_TYPE_TEXTS = {
    "regression": {"pt": "Regressão", "en": "regression"},
    "classification": {"pt": "Classificação", "en": "classification"},
    "clustering": {"pt": "Agrupamento", "en": "clustering"},
}


def _translate_task_type(task_type: str, lang: str) -> str:
    return _TASK_TYPE_TEXTS.get(task_type, {}).get(lang, task_type)


def _parse_job_started(job_name: str) -> Optional[datetime]:
    """Extracts the job's own creation date/time, embedded as the leading YYYY-MM-DD_HH-MM of
    its name (e.g. "2026-08-05_10-14_regression_RAUL_TRICHOMONAS" -> 2026-08-05 10:14) - the same
    value CONFIG's current_date field showed when the job was created. Returns None if job_name
    doesn't start with that pattern (e.g. a hand-renamed or otherwise unusual job folder)."""
    match = re.match(r"^(\d{4}-\d{2}-\d{2}_\d{2}-\d{2})", job_name or "")
    if not match:
        return None
    try:
        return datetime.strptime(match.group(1), "%Y-%m-%d_%H-%M")
    except ValueError:
        return None


def _format_working_time(delta, lang: str) -> str:
    """Formats a timedelta as "D Days : H Hours : M Minutes : S Seconds" (or the Portuguese
    equivalent), e.g. "2 Days : 4 Hours : 35 Minutes : 42 Seconds"."""
    total_seconds = max(0, int(delta.total_seconds()))
    days, remainder = divmod(total_seconds, 86400)
    hours, remainder = divmod(remainder, 3600)
    minutes, seconds = divmod(remainder, 60)
    texts = _METHOD_RESULTS_TEXTS
    return (
        f"{days} {texts['days'][lang]} : {hours} {texts['hours'][lang]} : "
        f"{minutes} {texts['minutes'][lang]} : {seconds} {texts['seconds'][lang]}"
    )


# --------------------------------------------------------------------------------------
# Entry point
# --------------------------------------------------------------------------------------

def generate_final_report(
    job_dir: str,
    job_name: str,
    state: Optional[dict[str, Any]] = None,
    progress_callback: Optional[Callable[[str], None]] = None,
    idioma: str = "pt",
    app_dir: Optional[str] = None,
) -> str:
    """Build <job_dir>/RESULTS/<job_name>_REPORT_<generation timestamp>.docx from the job's
    unified state JSON (job_dir/<job_name>.json) and the result files each STEP already writes to
    disk. The generation timestamp (not the job name's own date/time) means re-generating the
    report after changing parameters/re-running calculations never overwrites a previous report
    from the same job. Returns the output path. Raises RuntimeError if no section had any
    recorded state at all.

    idioma: UI language active when the report is generated ("en"/"pt", read from CODRUG's own
    self._idioma) - so far only affects the fixed Introduction section (see _add_intro_section)
    and the "Método e Resultados" header (title, Job Name/Task Type/Started/Finished/Working
    Time); the STEP 1-6 section bodies are still English-only, to be ported over topic by topic.
    app_dir: CODRUG's own install directory (CODRUG.py's self.dp_dir), used to find the fixed
    BASE/workflow.png asset for the Introduction section. Defaults to this module's own parent
    directory (MODULES/../) when not given, matching CODRUG.py's own dp_dir computation."""
    require_docx()
    app_dir = app_dir or os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

    def report(text: str) -> None:
        if progress_callback is not None:
            progress_callback(text)

    state = state if state is not None else load_job_settings(job_dir)
    if not isinstance(state, dict):
        state = {}

    document = Document()
    section = document.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    document.styles["Normal"].font.name = "Calibri"
    document.styles["Normal"].font.size = Pt(10)

    lang = "en" if idioma == "en" else "pt"
    now = datetime.now()  # reused below for "Finished"/Working Time and the output filename

    report("Building Introduction...")
    _add_intro_section(document, app_dir, idioma)

    header_texts = _METHOD_RESULTS_TEXTS
    _bar(document, header_texts["title"][lang], COLOR_TITLE_BAR)
    _field_line(document, header_texts["job_name"][lang], job_name)
    _field_line(document, header_texts["task_type"][lang], _translate_task_type(state.get("task_type", ""), lang))
    started_dt = _parse_job_started(job_name)
    if started_dt is not None:
        _field_line(document, header_texts["started"][lang], started_dt.strftime("%Y-%m-%d %H:%M"))
    _field_line(document, header_texts["finished"][lang], now.strftime("%Y-%m-%d %H:%M"))
    if started_dt is not None:
        _field_line(document, header_texts["working_time"][lang], _format_working_time(now - started_dt, lang))
    document.add_paragraph()

    report("Building STEP 1 - Dataset Preparation...")
    added_1 = _add_step1_section(document, job_dir, state, idioma)
    report("Building STEP 2 - Preprocessing and Exploratory Analysis...")
    added_23 = _add_step2_3_section(document, job_dir, state, idioma)
    report("Building STEP 3 - Features Engineering...")
    added_4 = _add_step4_section(document, job_dir, state, idioma)
    report("Building STEP 4 - Machine Learning Models...")
    added_5 = _add_step5_section(document, job_dir, state, idioma)
    report("Building STEP 5 - Applicability Domain and Similarity Analysis...")
    added_7 = _add_step7_section(document, job_dir, state, idioma)
    report("Building STEP 6 - Consensus Analysis...")
    added_8 = _add_step8_section(document, job_dir, state, idioma)

    if not any((added_1, added_23, added_4, added_5, added_7, added_8)):
        raise RuntimeError(
            "No recorded state was found for this job. Run at least one STEP (Generate Base "
            "Dataset, Outlier Elimination, Compute AD, Consensus Generate, etc.) before "
            "generating the final report."
        )

    results_dir = os.path.join(job_dir, "RESULTS")
    os.makedirs(results_dir, exist_ok=True)
    # Same "now" instant already shown as "Finished" above (not the job name's own date/time), so
    # re-generating the report after changing parameters/re-running calculations never overwrites
    # a previous report from the same job.
    output_path = os.path.join(results_dir, f"{job_name}_REPORT_{now.strftime('%Y-%m-%d_%H-%M')}.docx")
    report(f"Saving {output_path} ...")
    document.save(output_path)
    return output_path
